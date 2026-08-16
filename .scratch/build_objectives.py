# -*- coding: utf-8 -*-
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from objectives_data import OBJECTIVES, TYPE_NOTE
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol, w in zip(tblGrid.findall(qn('w:gridCol')), widths_cm):
            gridCol.set(qn('w:w'), str(int(w * 567)))


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_run(p, text, bold=False, size=10, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def fill_card_cell(cell, c):
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, f"{c['id']} — «{c['name']}»", bold=True, size=11)
    p2 = cell.add_paragraph()
    add_run(p2, f"{c['kind']}, {c['type']}", size=9, color="666666")
    if c.get("newcard"):
        p3 = cell.add_paragraph()
        add_run(p3, "НОВАЯ КАРТА (каталог 1.7.0, 15.08.2026)", size=8, color="B00000")
    if c.get("figure"):
        p4 = cell.add_paragraph()
        add_run(p4, "ЗАДАНИЕ-РИСУНОК (фигура из жетонов)", size=8, color="0060B0")


def fill_desc_cell(cell, c):
    cell.text = ""
    p0 = cell.paragraphs[0]
    add_run(p0, TYPE_NOTE.get(c["type"], ""), size=8, color="808080")

    p1 = cell.add_paragraph()
    add_run(p1, "Что делать (в двух словах): ", bold=True, size=10)
    add_run(p1, c["desc"], size=10)

    p2 = cell.add_paragraph()
    add_run(p2, "Обычное требование: ", bold=True, size=10)
    add_run(p2, c["req"], size=10)

    p3 = cell.add_paragraph()
    add_run(p3, "Усиленное требование: ", bold=True, size=10)
    add_run(p3, c["enh"], size=10)

    p4 = cell.add_paragraph()
    add_run(p4, "Награда за обычное: ", bold=True, size=10)
    add_run(p4, c["base"], size=10)

    p5 = cell.add_paragraph()
    add_run(p5, "+ доп. награда за усиленное: ", bold=True, size=10)
    add_run(p5, c["special"], size=10)

    p6 = cell.add_paragraph()
    add_run(p6, "Если сжечь вместо выполнения (утиль): ", bold=True, size=10)
    add_run(p6, c["top"], size=10)

    if c.get("note"):
        p7 = cell.add_paragraph()
        add_run(p7, "Замер: ", bold=True, size=9, color="B00000")
        add_run(p7, c["note"], size=9, color="B00000")


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    doc.add_heading("Карты заданий — как устроены сейчас", level=1)
    p = doc.add_paragraph()
    add_run(p, "Источник: ", bold=True)
    add_run(p, "data/cards/objectives.1.7.0.yaml (актуальная версия, подключена сводом 1.8.0) "
               "+ код kelium.engine.Predicates / kelium.cards.objectives.GroupNewWar. "
               "16.08.2026.")
    p2 = doc.add_paragraph()
    add_run(p2, "Как читать: ", bold=True)
    add_run(p2, "левая колонка — полное описание того, что карта делает СЕЙЧАС (условие, "
                "усиление, награды, чем можно сжечь вместо выполнения). Правая колонка — "
                "пусто, для ваших пометок: что оставить, что переделать, что выпилить.")

    all_cards = OBJECTIVES
    groups = [
        ("ВОЙНА И ДАВЛЕНИЕ (20 карт)", [c for c in all_cards if c["id"] in
            ["o03","o07","o11","o12","o14","o17","o21","o22","o23","o24","o25","o26","o27","o28","o29","o31","o41","o42","o43","o44"]]),
        ("ЭКОНОМИКА И НАУКА (17 карт)", [c for c in all_cards if c["id"] in
            ["o01","o02","o04","o05","o08","o13","o15","o16","o18","o19","o20","o30","o33","o36","o37","o38","o40"]]),
        ("ЗАДАНИЯ-РИСУНКИ (3 карты)", [c for c in all_cards if c.get("figure")]),
        ("НАЧАЛЬНЫЕ (8 карт, только раунд 1, без усиления и утили)", [c for c in all_cards if c["kind"]=="начальная"]),
    ]

    for title, cards in groups:
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Карта"
        hdr[1].text = "Как устроена сейчас"
        hdr[2].text = "Комментарий / правка"
        for hc in hdr:
            for pp in hc.paragraphs:
                for rr in pp.runs:
                    rr.bold = True
            shade_cell(hc, "D9D9D9")
        set_col_widths(table, [4.5, 15.5, 7.0])
        for c in cards:
            row = table.add_row().cells
            fill_card_cell(row[0], c)
            fill_desc_cell(row[1], c)
            row[2].text = ""
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Свои карты / эффекты, которых здесь нет", level=1)
    doc.add_paragraph("Место для новых идей — то, что вы хотите добавить и чего нет в текущем наборе.")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "Название / идея"
    hdr2[1].text = "Как должна работать"
    hdr2[2].text = "Заметка"
    for hc in hdr2:
        for pp in hc.paragraphs:
            for rr in pp.runs:
                rr.bold = True
        shade_cell(hc, "D9D9D9")
    set_col_widths(t2, [7.0, 13.0, 7.0])
    for _ in range(12):
        t2.add_row()

    out = os.path.join(HERE, "1_Карты_заданий.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    build()
