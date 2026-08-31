# -*- coding: utf-8 -*-
"""ДОКУМЕНТ ПО СУПЕР-ЗАДАНИЯМ 5.0 — «суперутиль или накопитель».

Формат 5.0 не похож на прежние супер-задания (нет ячеек, вскрытия и счётчика),
поэтому у него свой генератор: таблица двенадцати карт и по странице на карту —
имя, накопитель, суперутиль.

Запуск: python tools/документ-супер5.py [версия]
"""
import io
import sys

import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

ВЕРСИЯ = sys.argv[1] if len(sys.argv) > 1 else "5.0.0"
ИСТОЧНИК = f"data/cards/super_objectives.{ВЕРСИЯ}.yaml"
ЦЕЛЬ = f"docs/КАРТЫ СУПЕР ЗАДАНИЙ {ВЕРСИЯ}.docx"


def сер(п, текст, кегль=9):
    run = п.add_run(текст)
    run.font.size = Pt(кегль)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return run


def main():
    набор = yaml.safe_load(io.open(ИСТОЧНИК, encoding="utf-8"))
    карты = набор["super_objectives"]

    doc = Document()
    doc.add_heading(f"Супер-задания {ВЕРСИЯ} — «суперутиль или накопитель»", 0)
    п = doc.add_paragraph()
    сер(п, "Каждому игроку в подготовку втайне раздаётся ОДНА карта. Сыграть "
           "можно только одну половину: СУПЕРУТИЛЬ — разовый эффект, карта "
           "сжигается СПЕЦ-действием в любой момент; НАКОПИТЕЛЬ — победные "
           "очки в конце партии, если карта дожила нетронутой. Рубашки "
           "одинаковые: соперник знает, что у тебя что-то есть, но не знает "
           "что. Колонка «замечания» — под пометки дизайнера.", 10)

    doc.add_heading("Все карты одной таблицей", 1)
    т = doc.add_table(rows=1, cols=5)
    т.style = "Table Grid"
    т.alignment = WD_TABLE_ALIGNMENT.CENTER
    шапка = т.rows[0].cells
    for i, з in enumerate(["№", "карта", "НАКОПИТЕЛЬ (очки в конце)",
                           "СУПЕРУТИЛЬ (разовый эффект)", "замечания"]):
        шапка[i].paragraphs[0].add_run(з).bold = True
    ширины = [Cm(0.9), Cm(3.2), Cm(5.8), Cm(5.8), Cm(2.6)]
    for n, к in enumerate(карты, 1):
        ряд = т.add_row().cells
        ряд[0].text = str(n)
        ряд[1].text = str(к.get("name", к["id"]))
        ряд[2].text = str(к.get("stockpile", ""))
        ряд[3].text = str(к.get("burn", ""))
        ряд[4].text = ""
    for ряд in т.rows:
        for i, кл in enumerate(ряд.cells):
            кл.width = ширины[i]
            for пар in кл.paragraphs:
                for run in пар.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading("По одной карте", 1)
    for n, к in enumerate(карты, 1):
        doc.add_heading(f"{n}. «{к.get('name', к['id'])}»  ({к['id']})", 2)
        п = doc.add_paragraph()
        п.add_run("НАКОПИТЕЛЬ: ").bold = True
        п.add_run(str(к.get("stockpile", "")))
        п = doc.add_paragraph()
        п.add_run("СУПЕРУТИЛЬ: ").bold = True
        п.add_run(str(к.get("burn", "")))

    doc.save(ЦЕЛЬ)
    print(f"готово: {ЦЕЛЬ} карт: {len(карты)}")


if __name__ == "__main__":
    main()
