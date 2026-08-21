# -*- coding: utf-8 -*-
"""ДОКУМЕНТ ПО КАРТАМ ЗАДАНИЙ — собирается ИЗ САМИХ КАРТ.

Читает снимок каталога, выгруженный из кода (kelium.ВыгрузкаКаталога), и печатает
.docx: таблицу всех карт (номер, имя, условие, усиление, награды, утиль) и
разбор по темам с печатным текстом каждой карты.

Почему из снимка кода, а не из файла набора. У карты, живущей в коде, запись в
наборе — зеркало: условие там записано предикатом с параметрами, а человеку нужен
печатный текст. Снимок берёт ровно то, что карта печатает на себе.

Запуск:
    java kelium.ВыгрузкаКаталога objectives док "снимок для документа"
    python tools/документ-задания.py док 1.10.0
Второй аргумент — номер версии для заголовка документа.
"""
import io
import sys

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

СНИМОК = sys.argv[1] if len(sys.argv) > 1 else "док"
ВЕРСИЯ = sys.argv[2] if len(sys.argv) > 2 else "1.10.0"
ИСТОЧНИК = f"data/cards/objectives.{СНИМОК}.yaml"
ЦЕЛЬ = f"docs/КАРТЫ ЗАДАНИЙ {ВЕРСИЯ}.docx"

ПРИРОДА = {
    "state": "СОСТОЯНИЕ — проверяется в любой момент",
    "incident": "ПРОИСШЕСТВИЕ — должно случиться В ЭТОТ ХОД",
    "sacrifice": "ЖЕРТВА — карта требует что-то сдать",
}

НАГРАДЫ = {
    "coin": "монеты",
    "ammo": "боеприпасы",
    "debris": "обломки",
    "kelium": "келемий",
    "objective_card": "карты задания",
    "arsenal": "карта арсенала",
    "arsenal_from_display": "карта арсенала НА ВЫБОР из открытых",
    "module": "жетон модуля",
    "storage_token": "жетон хранилища",
    "vp": "победные очки",
}


def награда(узел):
    if not узел:
        return "—"
    части = []
    for ключ, значение in узел.items():
        имя = НАГРАДЫ.get(ключ, ключ)
        if ключ == "module":
            части.append("жетон модуля "
                         + ("атаки" if значение == "attack" else "сборки"))
        elif значение == 1:
            части.append(имя)
        else:
            части.append(f"{значение} {имя}")
    return ", ".join(части)


def условие(карта, ключ):
    узел = карта.get(ключ)
    if not isinstance(узел, dict):
        return "—"
    if "условие" in узел:
        return str(узел["условие"])
    # запись предикатом — покажем хотя бы имя и параметры, но это признак того,
    # что снимок сделан не из кода
    return str(узел)


def сер(п, текст, кегль=9):
    run = п.add_run(текст)
    run.font.size = Pt(кегль)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return run


def main():
    набор = yaml.safe_load(io.open(ИСТОЧНИК, encoding="utf-8"))
    карты = набор["objectives"]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.2)
    section.top_margin = section.bottom_margin = Cm(1.2)

    doc.add_heading(f"Карты заданий {ВЕРСИЯ}", level=0)
    п = doc.add_paragraph()
    сер(п, f"Кристаллы Раздора · карт в наборе: {len(карты)} · "
           "условия и тексты — как их печатает сама карта", 10)

    обычные = [c for c in карты if c.get("kind") != "starting"]
    начальные = [c for c in карты if c.get("kind") == "starting"]
    сВитрины = [c for c in карты
                if "arsenal_from_display" in str(c.get("base_reward"))
                + str(c.get("special_reward"))]
    сАрсеналом = [c for c in карты
                  if "arsenal" in str(c.get("base_reward"))
                  + str(c.get("special_reward"))]

    doc.add_heading("Что в этом наборе", level=1)
    for текст in [
        f"Обычных карт {len(обычные)}, начальных заданий {len(начальные)} "
        "(раздаются отдельным дополнением в первом раунде).",
        f"Карту арсенала в награду дают {len(сАрсеналом)} карт, из них "
        f"{len(сВитрины)} — картой НА ВЫБОР из открытых на витрине. Это правило "
        "от 21.08.2026: путь получения карт арсенала в игре был ровно один "
        "(обмен на планшете науки), и карт игроку приходило полторы за партию.",
        "Карта арсенала на выбор стоит на самых трудных условиях — это приз. "
        "Обычная карта арсенала стоит на заданиях, которые по замеру выполняются "
        "часто: награда, которую никто не получает, не даёт ничего.",
    ]:
        doc.add_paragraph(текст)

    doc.add_heading("Все карты набора", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    ширины = [Cm(1.3), Cm(3.4), Cm(8.4), Cm(5.4), Cm(4.0), Cm(3.5)]
    заголовки = ["№", "Название", "Условие", "Усиление",
                 "Награда", "Усиленная награда"]
    for i, (h, w) in enumerate(zip(заголовки, ширины)):
        cell = t.rows[0].cells[i]
        cell.width = w
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for c in карты:
        row = t.add_row()
        значения = [
            c["id"],
            c.get("name", ""),
            условие(c, "requirement"),
            условие(c, "enhanced"),
            награда(c.get("base_reward")),
            награда(c.get("special_reward")),
        ]
        for i, (v, w) in enumerate(zip(значения, ширины)):
            cell = row.cells[i]
            cell.width = w
            r = cell.paragraphs[0].add_run(str(v))
            r.font.size = Pt(8)
            if i == 0:
                r.bold = True

    doc.add_page_break()
    doc.add_heading("Карты по одной — с печатным текстом", level=1)
    for заголовок, группа in [("Обычные карты", обычные),
                              ("Начальные задания", начальные)]:
        if not группа:
            continue
        doc.add_heading(заголовок, level=2)
        for c in группа:
            п = doc.add_paragraph()
            r = п.add_run(f"{c['id']} · {c.get('name','')}")
            r.bold = True
            r.font.size = Pt(11)
            п = doc.add_paragraph()
            сер(п, ПРИРОДА.get(str(c.get("type")), str(c.get("type"))))
            п = doc.add_paragraph()
            сер(п, "УСЛОВИЕ: " + условие(c, "requirement"))
            if c.get("enhanced"):
                п = doc.add_paragraph()
                сер(п, "УСИЛЕНИЕ: " + условие(c, "enhanced"))
            п = doc.add_paragraph()
            сер(п, "НАГРАДА: " + награда(c.get("base_reward"))
                + "   |   УСИЛЕННАЯ: " + награда(c.get("special_reward")))
            верх = c.get("top") or {}
            if верх.get("label"):
                п = doc.add_paragraph()
                сер(п, "УТИЛЬ (сжечь ради этого): " + str(верх["label"]))
            if c.get("описание"):
                p = doc.add_paragraph(str(c["описание"]))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(10)

    doc.save(ЦЕЛЬ)
    print("готово:", ЦЕЛЬ, "карт:", len(карты))


if __name__ == "__main__":
    main()
