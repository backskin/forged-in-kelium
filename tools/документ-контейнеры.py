# -*- coding: utf-8 -*-
"""ДОКУМЕНТ ПО КОНТЕЙНЕРАМ — ПО ТИПАМ, а не по карточкам.

Почему по типам. У контейнеров 32 карты на 12 форм: перечислять их по одной
значит показывать одно и то же по пять раз. За столом важно другое — КАКИЕ
формы бывают и СКОЛЬКО каждой в колоде: именно из этого игрок считает свои
шансы при вскрытии.

Источник — снимок каталога, выгруженный ИЗ КОДА (kelium.ВыгрузкаКаталога):
контейнеры живут в классах, запись набора им зеркало.

Запуск:
    java kelium.ВыгрузкаКаталога containers доккон "снимок"
    python tools/документ-контейнеры.py доккон 4.0.0
"""
import collections
import io
import sys

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

СНИМОК = sys.argv[1] if len(sys.argv) > 1 else "доккон"
ВЕРСИЯ = sys.argv[2] if len(sys.argv) > 2 else "4.0.0"
ЦЕЛЬ = f"docs/КАРТЫ КОНТЕЙНЕРОВ {ВЕРСИЯ}.docx"

РЕСУРСЫ = {
    "ammo": "боеприпас",
    "coin": "монета",
    "trophy": "трофей",
    "kelium": "келемий",
    "objective_cards": "карта задания",
    "arsenal": "карта арсенала",
}

# КУРС ДЛЯ СВЕРКИ — тот же, что записан в самом наборе (см. шапку
# containers.4.0.0.yaml). Нужен, чтобы показать ценность формы одним числом и
# увидеть, ровно ли разложена колода.
КУРС = {
    "ammo": 1.0,
    "coin": 1.0,
    "trophy": 1.5,
    "kelium": 1.5,
    "objective_cards": 2.0,
    "arsenal": 2.5,
}

ТИРЫ = {"common": "обычный", "good": "хороший", "rare": "редкий"}


def форма(карта):
    """Что напечатано на карте — как человеческая строка и как ключ группировки."""
    a = карта.get("a") or {}
    if a.get("effect") == "empty":
        return "ПУСТО — ничего", 0.0
    части = []
    цена = 0.0
    for ключ, значение in (a.get("params") or {}).items():
        имя = РЕСУРСЫ.get(ключ, ключ)
        части.append(имя if значение == 1 else f"{значение} × {имя}")
        цена += КУРС.get(ключ, 0) * значение
    return " + ".join(части) if части else str(a), цена


def сер(п, текст, кегль=9):
    r = п.add_run(текст)
    r.font.size = Pt(кегль)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return r


def таблица(doc, заголовки, ширины, строки, жирнаяПервая=True):
    t = doc.add_table(rows=1, cols=len(заголовки))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(заголовки, ширины)):
        cell = t.rows[0].cells[i]
        cell.width = w
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for строка in строки:
        row = t.add_row()
        for i, (v, w) in enumerate(zip(строка, ширины)):
            cell = row.cells[i]
            cell.width = w
            r = cell.paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            if i == 0 and жирнаяПервая:
                r.bold = True
    return t


def main():
    карты = yaml.safe_load(
        io.open(f"data/cards/containers.{СНИМОК}.yaml", encoding="utf-8"))["containers"]

    # Группировка по форме: сохраняем порядок первого появления.
    группы = collections.OrderedDict()
    for c in карты:
        имя, цена = форма(c)
        g = группы.setdefault(имя, {"цена": цена, "карты": [], "тиры": collections.Counter()})
        g["карты"].append(c)
        g["тиры"][c.get("tier")] += 1

    всего = len(карты)

    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width
    s.left_margin = s.right_margin = Cm(1.2)
    s.top_margin = s.bottom_margin = Cm(1.2)

    doc.add_heading(f"Карты контейнеров {ВЕРСИЯ}", level=0)
    сер(doc.add_paragraph(),
        f"Кристаллы Раздора · карт в колоде: {всего} шт · форм: {len(группы)} шт · "
        "сгруппировано по тому, что карта выдаёт", 10)

    doc.add_heading("Как это работает", level=1)
    for т in [
        "Контейнеры НАПЕЧАТАНЫ на картонных блоках поля. Жетон, накрывший "
        "сектор с печатным контейнером, немедленно берёт карту контейнера из "
        "запаса рубашкой вверх. Вскрытие карты — свободное действие.",
        "ВЫБОРА НА КАРТЕ НЕТ: что напечатано, то и выдаётся сразу. Сторона «б» "
        "не заведена ни у одной карты — это решение версии 4.0, а не пропуск.",
        "Добыча берёт контейнер иначе: запитанный добытчик может забрать "
        "ОТКРЫТЫЙ печатный контейнер со своего или примыкающего гекса.",
        "Воздушный контейнер (по одному на каждой стороне блока) берёт только "
        "авиация.",
    ]:
        doc.add_paragraph(т)

    doc.add_heading("Формы карт и сколько их в колоде", level=1)
    сер(doc.add_paragraph(),
        "«Ценность» — по курсу сверки самого набора: монета 1, боеприпас 1, "
        "трофей 1.5, келемий 1.5, карта задания 2, карта арсенала 2.5.", 9)
    строки = []
    for имя, g in sorted(группы.items(), key=lambda kv: (-len(kv[1]["карты"]), kv[0])):
        n = len(g["карты"])
        тиры = ", ".join(f"{ТИРЫ.get(t, t)} {k}" for t, k in g["тиры"].most_common())
        строки.append([
            имя, n, f"{100 * n / всего:.0f} %", f"{g['цена']:.1f}", тиры,
            ", ".join(c["name"] for c in g["карты"]),
        ])
    таблица(doc,
            ["что выдаёт", "карт в колоде, шт", "доля колоды", "ценность по курсу",
             "тир", "названия карт"],
            [Cm(5.6), Cm(2.4), Cm(2.2), Cm(2.4), Cm(4.4), Cm(9.6)],
            строки)

    п = doc.add_paragraph()
    средняя = sum(g["цена"] * len(g["карты"]) for g in группы.values()) / всего
    сер(п, f"Средняя ценность карты колоды по курсу сверки: {средняя:.2f}.", 10)

    doc.add_heading("Разрез по ресурсам", level=1)
    сер(doc.add_paragraph(),
        "Сколько карт колоды выдают этот ресурс (карта с парой ресурсов "
        "попадает в обе строки) и сколько единиц ресурса лежит в колоде всего.", 9)
    поРесурсам = collections.Counter()
    единиц = collections.Counter()
    for c in карты:
        a = c.get("a") or {}
        for ключ, значение in (a.get("params") or {}).items():
            поРесурсам[ключ] += 1
            единиц[ключ] += значение
    строки2 = [[РЕСУРСЫ.get(k, k), поРесурсам[k], f"{100 * поРесурсам[k] / всего:.0f} %",
                единиц[k]] for k, _ in поРесурсам.most_common()]
    пусто = sum(1 for c in карты if (c.get("a") or {}).get("effect") == "empty")
    строки2.append(["ПУСТО (ничего)", пусто, f"{100 * пусто / всего:.0f} %", 0])
    таблица(doc, ["ресурс", "на скольких картах, шт", "доля колоды", "единиц в колоде, шт"],
            [Cm(5.0), Cm(4.0), Cm(3.0), Cm(4.4)], строки2)

    doc.add_heading("Разрез по тирам", level=1)
    тиры = collections.Counter(c.get("tier") for c in карты)
    таблица(doc, ["тир", "карт, шт", "доля колоды"],
            [Cm(4.0), Cm(3.0), Cm(3.0)],
            [[ТИРЫ.get(t, t), n, f"{100 * n / всего:.0f} %"] for t, n in тиры.most_common()])

    doc.add_page_break()
    doc.add_heading("Печатные тексты по формам", level=1)
    сер(doc.add_paragraph(),
        "Текст у карт одной формы разный — это разные предметы, а не копии. "
        "Здесь они собраны под своей формой.", 9)
    for имя, g in группы.items():
        п = doc.add_paragraph()
        r = п.add_run(f"{имя} — {len(g['карты'])} шт")
        r.bold = True
        r.font.size = Pt(11)
        for c in g["карты"]:
            сер(doc.add_paragraph(),
                f"{c['id']} · {c.get('name', '')} — {c.get('описание', '')}")

    doc.save(ЦЕЛЬ)
    print("готово:", ЦЕЛЬ, "| карт:", всего, "| форм:", len(группы))


if __name__ == "__main__":
    main()
