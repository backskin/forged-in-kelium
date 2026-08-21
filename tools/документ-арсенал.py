# -*- coding: utf-8 -*-
"""ДОКУМЕНТ ПО КАРТАМ АРСЕНАЛА — собирается ИЗ САМОГО НАБОРА.

Читает data/cards/arsenal.<версия>.yaml и печатает .docx: таблицу всех карт
(номер, имя, утиль, низ) и разбор новой редакции утиля.

Почему из данных, а не руками. Документ, набранный отдельно от набора, расходится
с игрой на первой же правке — и заметить это можно только за столом. Здесь
источник один: что в наборе, то и в документе.

Запуск: python tools/документ-арсенал.py [версия]
"""
import io
import sys

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ВЕРСИЯ = sys.argv[1] if len(sys.argv) > 1 else "3.0.0"
ИСТОЧНИК = f"data/cards/arsenal.{ВЕРСИЯ}.yaml"
ЦЕЛЬ = f"docs/КАРТЫ АРСЕНАЛА {ВЕРСИЯ}.docx"

# Как называется вид низа карты человеческим языком.
ВИД_НИЗА = {
    "POST": "постоянная способность",
    "SPEC": "СПЕЦ-действие",
    "SCORING": "очки в конце партии",
}

ГРУППЫ = [
    ("Скорость, прочность, щиты", ["b01", "b02", "b03", "b17", "b18", "b15", "b16",
                                    "b26", "b27", "b28"]),
    ("Производство и хранение", ["b04", "b05", "b06", "b07", "b19", "b20", "b21",
                                  "b13", "b25", "b11"]),
    ("Приказы и темп", ["b08", "b09", "b10", "b22", "b23", "b24", "b12", "b29"]),
    ("Бой и удар", ["b14", "b30", "b31", "b27"]),
    ("Стартовые карты (по одной каждому игроку)",
     ["bs1", "bs2", "bs3", "bs4", "bs5", "bs6", "bs7", "bs8"]),
    ("Карты-цели (очки в конце партии)", ["v01", "v02", "v03", "v04", "v05", "v06"]),
]


def сер(п, текст, кегль=9):
    """Серая пояснительная строка."""
    run = п.add_run(текст)
    run.font.size = Pt(кегль)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return run


def главное(doc, текст):
    п = doc.add_paragraph()
    run = п.add_run(текст)
    run.bold = True
    run.font.size = Pt(11)
    return п


def таблица_карт(doc, карты):
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    ширины = [Cm(1.6), Cm(4.2), Cm(7.6), Cm(11.6)]
    заголовки = ["№", "Название", "УТИЛЬ (верх)", "УСТАНОВКА (низ)"]
    for i, (h, w) in enumerate(zip(заголовки, ширины)):
        cell = t.rows[0].cells[i]
        cell.width = w
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for c in карты:
        row = t.add_row()
        значения = [
            c["id"],
            c.get("name", ""),
            c["top"].get("label", c["top"]["effect"]),
            (c.get("bottom") or {}).get("label", ""),
        ]
        for i, (v, w) in enumerate(zip(значения, ширины)):
            cell = row.cells[i]
            cell.width = w
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
    return t


def main():
    набор = yaml.safe_load(io.open(ИСТОЧНИК, encoding="utf-8"))
    карты = набор["arsenal"]
    по_номеру = {c["id"]: c for c in карты}

    doc = Document()
    section = doc.sections[0]
    # Альбомная ориентация: у карт четыре столбца текста, и в книжной они мнутся.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    doc.add_heading(f"Карты арсенала {ВЕРСИЯ}", level=0)
    п = doc.add_paragraph()
    сер(п, "Кристаллы Раздора · редакция утиля от 21.08.2026 · "
           f"карт в наборе: {len(карты)}", 10)

    doc.add_heading("Что изменилось и почему", level=1)
    for текст in [
        "НИЗ КАРТ НЕ ИЗМЕНЁН НИ У ОДНОЙ. Способности живут в коде, они проверены "
        "и работают; переделан только верх — утиль.",
        "ЗАМЕР, из которого всё выросло. Утиль «выполни действие» стоял на 15 "
        "картах из 45. Ровно эти карты боты не устанавливали НИКОГДА: за 300 "
        "партий 26 карт не были установлены ни разу. Причина не в ботах — "
        "действие в этой игре самый дорогой ресурс (их около двадцати четырёх на "
        "партию), и выбор «карта или лишнее действие» решался в пользу действия "
        "всегда.",
        "ПРАВИЛО НОВОГО УТИЛЯ: утиль должен быть сопоставим со своим низом, а не "
        "сильнее его. Поэтому целого бесплатного действия без оговорок в колоде "
        "больше нет: где действие уместно, у него есть предел («одна операция», "
        "«два жетона», «двумя добытчиками») либо оно идёт с надбавкой вместо "
        "широты («Стройка со скидкой», «Рынок обеими половинами»).",
        "ЧТО ПОЯВИЛОСЬ: плата за положение на поле («по 2 монеты за каждое своё "
        "войско», всегда с потолком), кража ресурса и кража карты арсенала, "
        "обновление витрины, десант на свободный гекс, обмен по печатной "
        "лестнице, золочение жетона модуля и шесть карт, закрывающих ВСЕ шесть "
        "пар щита из четырёх родов войск.",
    ]:
        doc.add_paragraph(текст)

    doc.add_heading("Чего в этой редакции нет", level=1)
    doc.add_paragraph(
        "Два утиля из списка требуют РЕАКЦИИ на чужое действие: «в чужой Бой "
        "забери инициативу и атакуй первым» и «после чужой Добычи забери у "
        "игрока 1 келемий». Реакций на чужой ход движок пока не умеет — это "
        "отдельный механизм, а не эффект. Карту, которая не работает, заводить "
        "нельзя, поэтому эти два пункта ждут механизма реакций.")

    doc.add_heading("Все карты набора", level=1)
    таблица_карт(doc, карты)

    doc.add_page_break()
    doc.add_heading("Карты по темам — с печатным текстом", level=1)
    напечатано = set()
    for имя, номера in ГРУППЫ:
        свои = [по_номеру[n] for n in номера
                if n in по_номеру and n not in напечатано]
        if not свои:
            continue
        напечатано.update(c["id"] for c in свои)
        doc.add_heading(имя, level=2)
        for c in свои:
            главное(doc, f"{c['id']} · {c.get('name','')}")
            низ = c.get("bottom") or {}
            п = doc.add_paragraph()
            сер(п, "УТИЛЬ: " + str(c["top"].get("label", c["top"]["effect"])))
            п = doc.add_paragraph()
            сер(п, ВИД_НИЗА.get(str(низ.get("kind")), "низ") + ": "
                   + str(низ.get("label", "")))
            текст = c.get("описание")
            if текст:
                p = doc.add_paragraph(str(текст))
                p.paragraph_format.space_after = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    остались = [c for c in карты if c["id"] not in напечатано]
    if остались:
        doc.add_heading("Остальные карты", level=2)
        for c in остались:
            главное(doc, f"{c['id']} · {c.get('name','')}")
            низ = c.get("bottom") or {}
            п = doc.add_paragraph()
            сер(п, "УТИЛЬ: " + str(c["top"].get("label", c["top"]["effect"])))
            п = doc.add_paragraph()
            сер(п, ВИД_НИЗА.get(str(низ.get("kind")), "низ") + ": "
                   + str(низ.get("label", "")))
            if c.get("описание"):
                doc.add_paragraph(str(c["описание"]))

    doc.save(ЦЕЛЬ)
    print("готово:", ЦЕЛЬ, "карт:", len(карты))


if __name__ == "__main__":
    main()
